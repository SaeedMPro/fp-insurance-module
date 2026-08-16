#!/usr/bin/env python3
"""
Builds the Persian project report as a Word (.docx) file.

Why a script and not a hand-made file: the report has to stay in sync with the
code (numbers, rule tables, acceptance criteria), and a script makes it
regenerable. Run it after any change worth reflecting in the deliverable:

    /tmp/claude-1000/docxvenv/bin/python docs/build_report_docx.py

Everything about the layout is deliberate for a Persian academic document:
right-to-left paragraphs and tables, justified body text, 1.5 line spacing,
A4 with generous margins for binding, and Tahoma — which every Windows/Office
install has, so the department's machine renders it exactly as intended
instead of substituting a font.
"""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = "Tahoma"          # universally available; correct Persian shaping
MONO_FONT = "Consolas"
BODY_SIZE = Pt(11)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)   # dark navy: prints well in greyscale


# --------------------------------------------------------------------------
# low-level RTL helpers
# --------------------------------------------------------------------------

def rtl_paragraph(p):
    """Mark a paragraph right-to-left (Word needs both bidi and the alignment)."""
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    return p


def set_run_font(run, name=BODY_FONT, size=BODY_SIZE, bold=False, color=INK):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    # cs = complex script (Arabic/Persian); without it Word may substitute.
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size.pt * 2)))
    rPr.append(szCs)
    if bold:
        bCs = OxmlElement("w:bCs")
        rPr.append(bCs)


def rtl_table(table):
    """Right-to-left table: first column renders on the right."""
    tblPr = table._tbl.tblPr
    bidiVisual = OxmlElement("w:bidiVisual")
    tblPr.append(bidiVisual)


def no_row_split(row):
    """Keep a table row whole. Without this a row can break across a page
    boundary and leave what looks like an empty stray row on the next page."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    """Repeat this row at the top of each page the table continues onto."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


# --------------------------------------------------------------------------
# content builders
# --------------------------------------------------------------------------

def para(doc, text="", size=BODY_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=8, space_before=6, color=INK, line=1.5, font=BODY_FONT):
    p = doc.add_paragraph()
    rtl_paragraph(p)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if text:
        set_run_font(p.add_run(text), name=font, size=size, bold=bold, color=color)
    return p


def rich(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, line=1.5):
    """A paragraph made of (text, bold) fragments — for inline emphasis."""
    p = doc.add_paragraph()
    rtl_paragraph(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    for text, bold in parts:
        set_run_font(p.add_run(text), bold=bold)
    return p


def heading(doc, text, level=1, new_page=False):
    """A section heading. new_page starts it on a fresh page via
    page_break_before — inserting a separate break paragraph instead would
    leave a blank page whenever the previous block happened to end at the
    bottom of a page."""
    sizes = {1: Pt(15), 2: Pt(12.5), 3: Pt(11.5)}
    before = {1: 18, 2: 12, 3: 8}
    p = doc.add_paragraph()
    rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    if new_page:
        pf.page_break_before = True
    pf.space_before = Pt(0 if new_page else before[level])
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    set_run_font(p.add_run(text), size=sizes[level], bold=True,
                 color=ACCENT if level <= 2 else INK)
    # Word's built-in outline level, so the automatic table of contents finds it.
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level - 1))
    pPr.append(outline)
    return p


def bullets(doc, items, numbered=False):
    for i, item in enumerate(items, 1):
        marker = f"{to_fa(i)}. " if numbered else "•  "
        p = doc.add_paragraph()
        rtl_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = 1.4
        pf.right_indent = Cm(0.6)
        if isinstance(item, tuple):
            set_run_font(p.add_run(marker))
            set_run_font(p.add_run(item[0]), bold=True)
            set_run_font(p.add_run(item[1]))
        else:
            set_run_font(p.add_run(marker + item))


def table(doc, headers, rows, widths=None, header_fill="1F3A5F", zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Fixed layout: without this Word/LibreOffice auto-fit ignores the column
    # widths below and rebalances columns by content.
    t.autofit = False
    rtl_table(t)

    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        rtl_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        set_run_font(p.add_run(h), size=Pt(10), bold=True,
                     color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(cell, header_fill)

    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for c_i, value in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            rtl_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            bold = value.startswith("**") and value.endswith("**")
            set_run_font(p.add_run(value.strip("*")), size=Pt(10), bold=bold)
            if zebra and r_i % 2 == 1:
                shade(cells[c_i], "F2F5F9")

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        t._tbl.tblPr.append(layout)
        # With a fixed layout the renderer sizes columns from tblGrid, not from
        # the cell widths, so the grid has to be updated too or the columns
        # stay evenly divided.
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for col, w in zip(grid.findall(qn("w:gridCol")), widths):
                col.set(qn("w:w"), str(int(Cm(w).twips)))

    for row in t.rows:
        no_row_split(row)
    repeat_header(t.rows[0])

    # No spacer paragraph after the table: an empty paragraph here gets pushed
    # onto the next page whenever a table ends at the page bottom, and if the
    # following heading also starts a new page the result is a blank page.
    # The gap comes from the following block's space_before instead.
    return t


def code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.15
        pf.left_indent = Cm(0.8)
        set_run_font(p.add_run(line), name=MONO_FONT, size=Pt(9.5), color=RGBColor(0x22, 0x22, 0x22))
    para(doc, "", space_after=6)


def note(doc, text):
    """An indented, muted remark — used for caveats and asides."""
    p = doc.add_paragraph()
    rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.right_indent = Cm(0.8)
    pf.left_indent = Cm(0.8)
    pf.space_before = Pt(4)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.4
    set_run_font(p.add_run(text), size=Pt(10), color=MUTED)


FA_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")


def to_fa(value):
    return str(value).translate(FA_DIGITS)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_numbers(section):
    """Footer: 'صفحه N' centred, via a Word PAGE field."""
    footer = section.footer
    p = footer.paragraphs[0]
    rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("صفحه "), size=Pt(9), color=MUTED)
    run = p.add_run()
    set_run_font(run, size=Pt(9), color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._element.addnext(fld)


def add_toc(doc):
    """Insert a real table-of-contents field; Word fills it on open/F9."""
    p = doc.add_paragraph()
    rtl_paragraph(p)
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "برای نمایش فهرست، روی آن کلیک کنید و کلید F9 را بزنید."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (fldChar, instr, sep, placeholder, end):
        run._element.append(el)


# --------------------------------------------------------------------------
# the document
# --------------------------------------------------------------------------

def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:cs"), BODY_FONT)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(3.0)   # wider gutter on the right for binding (RTL)
    # Make the whole section RTL so page numbering/columns follow Persian order.
    sectPr = sec._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)

    # ---------------------------------------------------------------- cover
    para(doc, "", space_after=40)
    para(doc, "دانشگاه بوعلی سینا", size=Pt(13), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, color=ACCENT)
    para(doc, "دانشکده مهندسی — گروه مهندسی کامپیوتر", size=Pt(11.5),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=70, color=MUTED)

    para(doc, "گزارش پروژه پایانی دوره کارشناسی", size=Pt(11.5),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16, color=MUTED)
    para(doc, "طراحی و پیاده‌سازی ماژول بیمه تکمیلی", size=Pt(21), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, color=ACCENT)
    para(doc, "سامانه تحت وب با موتور قانون داده‌محور، گردش کار چندمرحله‌ای و ردیابی کامل",
         size=Pt(11.5), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=80, color=MUTED)

    # Cover details as plain centred lines — a bordered table here would look
    # like part of the report body.
    for label, value in [
        ("نگارنده", "سعید مظاهری"),
        ("شماره دانشجویی", "۴۰۲۴۲۹۹۹۰۰۴"),
        ("رشته و گرایش", "مهندسی کامپیوتر — نرم‌افزار"),
        ("استاد راهنما", "جناب آقای دکتر مهدی سخایی‌نیا"),
        ("مدت پروژه", "نیم‌سال ۴۰۴۲ و ۴۰۴۳"),
    ]:
        p = doc.add_paragraph()
        rtl_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.2
        set_run_font(p.add_run(f"{label}:  "), size=Pt(11), color=MUTED)
        set_run_font(p.add_run(value), size=Pt(12), bold=True)

    page_break(doc)
    add_page_numbers(sec)
    # No page number on the cover.
    sec.different_first_page_header_footer = True

    # ------------------------------------------------------------- abstract
    heading(doc, "چکیده", 1, new_page=True)
    para(doc,
         "در بسیاری از سازمان‌ها، فرایند بیمه تکمیلی کارکنان به‌صورت دستی و با تکیه بر "
         "صفحه‌های گسترده اداره می‌شود. این روش سه اشکال جدی دارد: نخست، محاسبه دستی سهم "
         "قابل پرداخت — که همزمان به درصد پوشش، سقف هر دفعه و مانده سقف سالانه وابسته است — "
         "مستعد خطاست؛ دوم، تصمیم‌ها قابل ردیابی نیستند و در صورت اعتراض کارمند، مبنای "
         "تصمیم قابل بازسازی نیست؛ سوم، هر تغییر در سیاست‌های رفاهی سازمان مستلزم تغییر در "
         "منطق محاسبه و در نتیجه نیازمند چرخه توسعه نرم‌افزار است.")
    para(doc,
         "در این پروژه یک سامانه تحت وب برای مدیریت بیمه تکمیلی طراحی و پیاده‌سازی کرده‌ام "
         "که این سه مسئله را هدف گرفته است. دستاورد اصلی، «موتور قانون پوشش و محاسبه سقف "
         "تعهدات» است که به‌جای کدنویسی ثابت، بر پایه پیکربندی داده‌محور کار می‌کند: درصد "
         "پوشش، سقف هر دفعه، سقف سالانه، دوره انتظار و نسبت‌های مجاز، همه به‌صورت رکوردهای "
         "نسخه‌دار در پایگاه داده نگه‌داری می‌شوند. در نتیجه تغییر سیاست رفاهی صرفاً از طریق "
         "پیکربندی و بدون تغییر کد و استقرار مجدد اعمال می‌شود.")
    para(doc,
         "در کنار آن، گردش کار چندمرحله‌ای درخواست با مسیرهای تأیید، رد و بازگشت برای تکمیل "
         "مدارک پیاده شده است و تمام رویدادهای مؤثر — همراه با کاربر، زمان و مقدار قبل و بعد — "
         "در یک لاگ ممیزی ثبت می‌شوند؛ به‌گونه‌ای که نوشتن لاگ و تغییر داده در یک تراکنش "
         "انجام می‌شود و از نظر ساختاری امکان تغییر بدون ردپا وجود ندارد. کنترل دسترسی مبتنی "
         "بر نقش در چهار سطح، گزارش‌گیری مدیریتی، و رابط کاربری کامل فارسی و راست‌به‌چپ نیز "
         "پیاده‌سازی شده‌اند.")
    para(doc,
         "سامانه با زبان Go در سمت سرور، پایگاه داده PostgreSQL و کتابخانه React با "
         "TypeScript در سمت کاربر توسعه یافته و درستی آن با آزمون‌های واحد، آزمون‌های "
         "یکپارچگی روی پایگاه داده واقعی و یک مجموعه آزمون سرتاسری روی مرورگر راستی‌آزمایی "
         "شده است.")
    para(doc, "", space_after=6)
    rich(doc, [("کلمات کلیدی: ", True),
               ("بیمه تکمیلی، سامانه تحت وب، موتور قانون، گردش کار چندمرحله‌ای، "
                "ردیابی و ممیزی، کنترل دسترسی مبتنی بر نقش، معماری لایه‌ای، Go، PostgreSQL", False)])

    # ----------------------------------------------------------------- TOC
    heading(doc, "فهرست مطالب", 1, new_page=True)
    note(doc, "این فهرست به‌صورت خودکار ساخته می‌شود. اگر خالی نمایش داده شد، روی آن کلیک "
              "کنید و کلید F9 را بزنید (یا در Word: References → Update Table).")
    add_toc(doc)

    # ------------------------------------------------------- 1. introduction
    heading(doc, "۱. مقدمه و بیان مسئله", 1, new_page=True)
    para(doc,
         "بیمه تکمیلی درمان یکی از رایج‌ترین خدمات رفاهی سازمان‌ها است. در الگوی متداول، "
         "کارمند فاکتور هزینه درمانی خود را به واحد رفاه تحویل می‌دهد، کارشناس آن را بررسی "
         "و سهم قابل پرداخت را محاسبه می‌کند و نتیجه برای پرداخت به واحد مالی اعلام می‌شود. "
         "هرچند این فرایند در ظاهر ساده است، در عمل چند دشواری مشخص دارد که انگیزه اصلی "
         "این پروژه بوده‌اند.")

    heading(doc, "۱.۱. دشواری محاسبه", 2)
    para(doc,
         "مبلغ قابل پرداخت تابعی از چند پارامتر همزمان است: درصد پوشش نوع خدمت، سقف مجاز "
         "برای هر دفعه، و مهم‌تر از همه مانده سقف سالانه که به سابقه مصرف همان فرد در همان "
         "سال قراردادی بستگی دارد. محاسبه دستی این ترکیب، به‌ویژه هنگام نزدیک شدن به سقف "
         "سالانه، مستعد خطاست و خطای آن مستقیماً بر حقوق مالی کارمند اثر می‌گذارد.")

    heading(doc, "۱.۲. نبود شفافیت و ردیابی", 2)
    para(doc,
         "در روش دستی، کارمند از وضعیت درخواست خود و از مانده سهمیه‌اش بی‌اطلاع است و برای "
         "هر پرسش باید به واحد رفاه مراجعه کند. مهم‌تر آن‌که اگر ماه‌ها بعد اعتراضی مطرح شود، "
         "مستندی وجود ندارد که نشان دهد آن تصمیم را چه کسی، در چه زمانی و بر مبنای کدام "
         "قاعده گرفته است.")

    heading(doc, "۱.۳. هزینه تغییر سیاست", 2)
    para(doc,
         "سیاست‌های رفاهی سازمان‌ها ثابت نیستند؛ درصد پوشش‌ها، سقف‌ها و شرایط واجدبودن هر "
         "سال بازنگری می‌شوند. اگر این قواعد در کد نرم‌افزار نوشته شده باشند، هر تغییر "
         "کوچک به یک چرخه کامل توسعه، آزمون و استقرار تبدیل می‌شود. این وابستگی، سازمان را "
         "برای اجرای تصمیم‌های خودش به تیم فنی وابسته می‌کند.")

    heading(doc, "۱.۴. رویکرد این پروژه", 2)
    para(doc,
         "بر پایه این سه مسئله، سامانه‌ای طراحی کرده‌ام که محاسبه را خودکار و یکسان می‌کند، "
         "هر رویداد را قابل ردیابی نگه می‌دارد، و — به‌عنوان تصمیم محوری طراحی — قواعد بیمه "
         "را از کد جدا کرده و به داده تبدیل می‌کند. در ادامه این گزارش، ابتدا اهداف و "
         "محدوده پروژه را مشخص می‌کنم، سپس تحلیل و طراحی، معماری و پیاده‌سازی را شرح می‌دهم، "
         "و در پایان نتایج را در برابر معیارهای پذیرشی که در پیشنهاد پروژه تعریف شده بود "
         "ارزیابی می‌کنم.")

    # ------------------------------------------------------------ 2. goals
    heading(doc, "۲. اهداف پروژه", 1)
    para(doc, "اهداف پروژه را در پنج بند صورت‌بندی کرده‌ام:")
    bullets(doc, [
        ("محاسبه خودکار و دقیق تعهدات: ",
         "تعیین مبلغ قابل پرداخت و مانده سقف برای هر فرد، بدون دخالت محاسبه دستی."),
        ("پیکربندی‌پذیری قواعد: ",
         "امکان تغییر درصد پوشش، سقف‌ها، دوره انتظار و شرایط واجدبودن صرفاً از طریق "
         "پیکربندی و بدون تغییر کد."),
        ("گردش کار شفاف: ",
         "پیاده‌سازی فرایند چندمرحله‌ای درخواست با مسیرهای تأیید، رد با ذکر دلیل، و "
         "بازگشت برای تکمیل مدارک."),
        ("ردیابی و پاسخ‌گویی: ",
         "ثبت تمام رویدادهای مؤثر همراه با کاربر، زمان و مقدار قبل و بعد، و امکان "
         "بازسازی تاریخچه هر درخواست."),
        ("گزارش‌گیری مدیریتی: ",
         "ارائه مصرف و مانده هر فرد، هزینه کل سازمان، و تفکیک بر اساس نوع خدمت و بازه زمانی."),
    ])
    para(doc,
         "علاوه بر این، از آن‌جا که این ماژول بخشی از «سامانه جامع خدمات رفاهی کارکنان» "
         "در نظر گرفته شده است، آن را با مرز مشخص و رابط برنامه‌نویسی تعریف‌شده طراحی "
         "کرده‌ام تا در آینده در کنار ماژول‌های دیگر (وام، مهمان‌سرا، خانه سازمانی) قابل "
         "استقرار باشد.")

    # ------------------------------------------------------------ 3. scope
    heading(doc, "۳. محدوده پروژه", 1)
    heading(doc, "۳.۱. داخل محدوده", 2)
    bullets(doc, [
        "تعریف قرارداد بیمه، طرح‌های پوشش و انواع خدمت",
        "تعریف قواعد پوشش به‌صورت نسخه‌دار و قابل پیکربندی",
        "ثبت کارکنان و اعضای تحت تکفل و اتصال هر کارمند به یک طرح",
        "ثبت درخواست هزینه برای شخص بیمه‌شده یا اعضای تحت تکفل",
        "محاسبه خودکار مبلغ قابل پرداخت و مانده سقف",
        "گردش کار چندمرحله‌ای با شاخه‌های تأیید، رد و بازگشت",
        "بارگذاری و دریافت مدارک پیوست درخواست، محدود به وضعیت‌های مجاز",
        "لاگ ممیزی کامل با امکان جست‌وجو و بازسازی تاریخچه",
        "کنترل دسترسی مبتنی بر نقش در چهار سطح",
        "گزارش‌گیری مدیریتی و رابط کاربری کامل فارسی",
        "رابط برنامه‌نویسی برای اتصال به سامانه مادر",
    ])
    heading(doc, "۳.۲. خارج از محدوده", 2)
    para(doc,
         "موارد زیر آگاهانه و از ابتدا خارج از محدوده تعریف شده‌اند و در این گزارش نیز "
         "به‌عنوان محدودیت شناخته‌شده معرفی می‌شوند، نه نقص:")
    table(doc,
          ["مورد", "توضیح"],
          [
              ["درگاه پرداخت واقعی", "پرداخت به‌صورت شبیه‌سازی ثبت می‌شود و شماره پیگیری تولید می‌گردد؛ اتصال به سامانه بانکی انجام نشده است."],
              ["اتصال زنده به سامانه‌های سازمان", "به‌جای اتصال زنده به منابع انسانی و مالی، یک درگاه برنامه‌نویسی با احراز هویت کلید ساخته شده است."],
              ["سایر ماژول‌های رفاهی", "وام، مهمان‌سرا و خانه سازمانی خارج از محدوده این ماژول‌اند."],
              ["آزمایشگاه", "خارج از محدوده."],
          ],
          widths=[4.5, 10.5])

    # ------------------------------------------------- 4. analysis & design
    heading(doc, "۴. تحلیل و طراحی", 1, new_page=True)

    heading(doc, "۴.۱. نقش‌ها و سطوح دسترسی", 2)
    para(doc,
         "چهار نقش تعریف شده است و هر نقش تنها به بخشی از سامانه دسترسی دارد که برای "
         "انجام وظیفه‌اش لازم است. جدول زیر مهم‌ترین قابلیت‌ها را به تفکیک نقش نشان می‌دهد.")
    table(doc,
          ["قابلیت", "کارمند", "کارشناس", "مدیر", "ممیز"],
          [
              ["مشاهده درخواست‌های خود", "دارد", "دارد", "دارد", "دارد"],
              ["مشاهده درخواست‌های دیگران", "ندارد", "دارد", "دارد", "دارد"],
              ["ثبت درخواست", "برای خود", "ندارد", "برای همه", "ندارد"],
              ["تأیید / رد / بازگرداندن", "ندارد", "دارد", "دارد", "ندارد"],
              ["ثبت پرداخت و بستن پرونده", "ندارد", "دارد", "دارد", "ندارد"],
              ["مشاهده مانده سقف", "خود", "همه", "همه", "ندارد"],
              ["مدیریت کارکنان و تحت تکفل", "ندارد", "فقط مشاهده", "دارد", "ندارد"],
              ["**انتشار نسخه جدید قانون**", "ندارد", "ندارد", "**دارد**", "ندارد"],
              ["مدیریت کاربران", "ندارد", "ندارد", "دارد", "ندارد"],
              ["گزارش‌ها و لاگ ممیزی", "ندارد", "ندارد", "دارد", "دارد"],
          ],
          widths=[6.2, 2.2, 2.2, 2.2, 2.2])
    para(doc,
         "کنترل دسترسی در دو لایه اعمال می‌شود: در لایه مسیرهای وب، نقش کاربر بررسی "
         "می‌شود؛ و در لایه سرویس، مالکیت داده کنترل می‌گردد. به این ترتیب یک کارمند حتی "
         "با دانستن شناسه درخواست دیگری، به آن دسترسی پیدا نمی‌کند.")

    heading(doc, "۴.۲. موجودیت‌های اصلی داده", 2)
    para(doc, "پایگاه داده از دوازده جدول تشکیل شده است. موجودیت‌های اصلی و نقش هرکدام:")
    table(doc,
          ["موجودیت", "نقش در سامانه"],
          [
              ["قرارداد بیمه", "بالاترین سطح؛ دارای تاریخ شروع و پایان"],
              ["طرح پوشش", "سطح خدمات زیر یک قرارداد؛ هر کارمند به یک طرح متصل است"],
              ["نوع خدمت", "دسته هزینه: ویزیت، دارو، دندان‌پزشکی، بستری، عینک"],
              ["**قانون پوشش**", "**پارامترهای محاسبه برای هر ترکیب طرح و نوع خدمت؛ نسخه‌دار**"],
              ["کارمند", "اطلاعات پرسنلی، تاریخ استخدام و وضعیت اشتغال"],
              ["عضو تحت تکفل", "همسر، فرزند یا والدین وابسته به کارمند"],
              ["کاربر", "حساب دسترسی با نقش؛ در صورت نقش کارمند به پرونده پرسنلی متصل"],
              ["درخواست", "هزینه ثبت‌شده با وضعیت و مبالغ محاسبه‌شده"],
              ["پرداخت", "رکورد پرداخت شبیه‌سازی‌شده با شماره پیگیری"],
              ["لاگ ممیزی", "رویدادها با کاربر، زمان و مقدار قبل و بعد"],
              ["کلید یکپارچه‌سازی", "کلید دسترسی سامانه مادر (به‌صورت هش‌شده)"],
              ["مدارک پیوست", "جدول پیش‌بینی‌شده برای پیوست فاکتور"],
          ],
          widths=[4.5, 10.5])

    heading(doc, "۴.۳. گردش کار درخواست", 2)
    para(doc,
         "گردش کار به‌صورت یک ماشین حالت صریح پیاده شده است. هر گذار مجاز در یک جدول "
         "تعریف شده و هر گذاری که در آن جدول نباشد رد می‌شود؛ بنابراین پریدن از مراحل "
         "امکان‌پذیر نیست.")
    # A monospace ASCII diagram breaks badly once Persian text is inside it,
    # so the three paths are written as plain right-to-left lines instead.
    for label, path in [
        ("مسیر اصلی", "پیش‌نویس ← ثبت‌شده ← در حال بررسی ← تأییدشده ← پرداخت‌شده ← بسته‌شده"),
        ("مسیر رد", "در حال بررسی ← رد شده ← بسته‌شده"),
        ("مسیر تکمیل مدارک", "در حال بررسی ← بازگشت برای تکمیل مدارک ← ثبت‌شده (ارسال مجدد)"),
    ]:
        p_ = doc.add_paragraph()
        rtl_paragraph(p_)
        p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_.paragraph_format.space_after = Pt(4)
        p_.paragraph_format.line_spacing = 1.3
        set_run_font(p_.add_run(f"{label}:  "), size=Pt(10), color=MUTED)
        set_run_font(p_.add_run(path), size=Pt(10.5), bold=True)
    para(doc, "", space_after=6)
    table(doc,
          ["وضعیت مبدأ", "اقدام", "وضعیت مقصد", "مجاز برای"],
          [
              ["پیش‌نویس", "ثبت", "ثبت‌شده", "مالک یا مدیر"],
              ["ثبت‌شده", "شروع بررسی", "در حال بررسی", "کارشناس یا مدیر"],
              ["در حال بررسی", "تأیید", "تأییدشده (با محاسبه مبلغ)", "کارشناس یا مدیر"],
              ["در حال بررسی", "رد (با ذکر دلیل)", "رد شده", "کارشناس یا مدیر"],
              ["در حال بررسی", "بازگرداندن (با ذکر دلیل)", "بازگشت برای تکمیل مدارک", "کارشناس یا مدیر"],
              ["بازگشت برای مدارک", "ارسال مجدد", "ثبت‌شده", "مالک یا مدیر"],
              ["تأییدشده", "ثبت پرداخت", "پرداخت‌شده", "کارشناس یا مدیر"],
              ["پرداخت‌شده یا رد شده", "بستن", "بسته‌شده", "کارشناس یا مدیر"],
          ],
          widths=[3.6, 3.4, 4.4, 3.6])
    para(doc,
         "مسیر «بازگشت برای تکمیل مدارک» تنها زمانی معنا دارد که کارمند بتواند مدرکِ "
         "خواسته‌شده را اضافه کند؛ به همین دلیل بارگذاری مدارک بخشی از همین ماشین حالت است "
         "و نه یک قابلیت جانبی. افزودن مدرک تنها در دو وضعیت «پیش‌نویس» و «بازگشت برای "
         "تکمیل مدارک» مجاز است. با ارسال درخواست، مدارک قفل می‌شوند: اگر پس از ارسال "
         "امکان تغییر مدرک وجود داشت، مستنداتی که کارشناس بر پایه آن‌ها تصمیم گرفته است "
         "می‌توانست بعد از تصمیم عوض شود و ارزش بررسی از میان می‌رفت.")

    heading(doc, "۴.۴. موتور قانون و محاسبه سقف", 2)
    para(doc,
         "این بخش، دستاورد مهندسی اصلی پروژه است. به‌جای آن‌که قواعدی مانند «پوشش "
         "دندان‌پزشکی پنجاه درصد است» در کد نوشته شود، هر قانون یک رکورد در پایگاه داده "
         "است با این پارامترها: درصد پوشش، سقف هر دفعه، سقف سالانه، دوره انتظار، و فهرست "
         "نسبت‌های مجاز. جدول زیر قواعد پیش‌فرض طرح استاندارد را نشان می‌دهد.")
    table(doc,
          ["نوع خدمت", "درصد پوشش", "سقف هر دفعه (ریال)", "سقف سالانه (ریال)", "دوره انتظار", "نسبت‌های مجاز"],
          [
              ["ویزیت", "۷۰٪", "۵۰۰٫۰۰۰", "۵٫۰۰۰٫۰۰۰", "۰ روز", "خود، همسر، فرزند، والدین"],
              ["دارو", "۸۰٪", "۱٫۰۰۰٫۰۰۰", "۱۰٫۰۰۰٫۰۰۰", "۰ روز", "خود، همسر، فرزند، والدین"],
              ["دندان‌پزشکی", "۵۰٪", "۳٫۰۰۰٫۰۰۰", "۱۵٫۰۰۰٫۰۰۰", "۹۰ روز", "خود، همسر، فرزند"],
              ["بستری", "۹۰٪", "۵۰٫۰۰۰٫۰۰۰", "۱۰۰٫۰۰۰٫۰۰۰", "۳۰ روز", "خود، همسر، فرزند، والدین"],
              ["عینک", "۶۰٪", "۲٫۰۰۰٫۰۰۰", "۴٫۰۰۰٫۰۰۰", "۱۸۰ روز", "خود، همسر، فرزند"],
          ],
          widths=[2.6, 1.9, 2.9, 2.9, 1.9, 2.8])

    para(doc, "هنگام تأیید یک درخواست، موتور قانون ابتدا پنج پیش‌شرط را بررسی می‌کند:")
    bullets(doc, [
        "کارمند در وضعیت شاغل باشد.",
        "اگر ذی‌نفع عضو تحت تکفل است، واقعاً به همان کارمند تعلق داشته باشد.",
        "برای ترکیب طرح و نوع خدمت، در تاریخ فاکتور قانون فعالی وجود داشته باشد.",
        "نسبت ذی‌نفع در فهرست نسبت‌های مجاز آن قانون باشد.",
        "دوره انتظار از تاریخ استخدام سپری شده باشد.",
    ], numbered=True)
    para(doc, "در صورت گذر از این پیش‌شرط‌ها، محاسبه در سه گام انجام می‌شود:")
    bullets(doc, [
        "مبلغ پوشش‌داده‌شده برابر است با مبلغ درخواستی ضرب در درصد پوشش، گرد شده به ریال کامل.",
        "اگر این مبلغ از سقف هر دفعه بیشتر باشد، به سقف هر دفعه محدود می‌شود.",
        "اگر همچنان از مانده سقف سالانه بیشتر باشد، به مانده سالانه محدود می‌شود.",
    ], numbered=True)
    para(doc,
         "برای روشن‌شدن اثر سقف‌ها، دو نمونه محاسبه برای خدمت ویزیت در طرح استاندارد "
         "(هفتاد درصد پوشش، سقف هر دفعه پانصد هزار ریال):")
    table(doc,
          ["مبلغ درخواستی", "درصد پوشش", "پیش از اعمال سقف", "مبلغ قابل پرداخت", "قید مؤثر"],
          [
              ["۳۵۰٫۰۰۰", "۷۰٪", "۲۴۵٫۰۰۰", "**۲۴۵٫۰۰۰**", "زیر همه سقف‌ها"],
              ["۱٫۰۰۰٫۰۰۰", "۷۰٪", "۷۰۰٫۰۰۰", "**۵۰۰٫۰۰۰**", "سقف هر دفعه"],
          ],
          widths=[3.0, 2.2, 3.2, 3.2, 3.4])

    heading(doc, "۴.۵. نسخه‌دار بودن قواعد", 2)
    para(doc,
         "قواعد هرگز در جای خود ویرایش نمی‌شوند. هنگام تغییر سیاست، نسخه فعلی با ثبت "
         "«تاریخ اعمال تا» بسته می‌شود و نسخه جدید درج می‌گردد. انتخاب قانون برای هر "
         "درخواست نیز بر اساس تاریخ فاکتور آن درخواست انجام می‌شود، نه تاریخ امروز.")
    note(doc,
         "اهمیت این تصمیم در ممیزی روشن می‌شود: اگر قواعد تاریخ‌دار نبودند، هر تغییر "
         "سیاست، محاسبه درخواست‌های گذشته را نیز تغییر می‌داد و بازسازی مبنای تصمیم‌های "
         "پیشین ناممکن می‌شد. همچنین برای حالت خاصِ انتشار دو نسخه در یک روز، تاریخ بستن "
         "نسخه پیشین به تاریخ شروع خودش محدود می‌شود و در انتخاب قانون فعال، نسخه جدیدتر "
         "اولویت دارد؛ برای این حالت آزمون اختصاصی نوشته‌ام.")

    # -------------------------------------------------------- 5. architecture
    heading(doc, "۵. معماری نرم‌افزار", 1, new_page=True)
    heading(doc, "۵.۱. فناوری‌های استفاده‌شده", 2)
    table(doc,
          ["بخش", "فناوری", "دلیل انتخاب"],
          [
              ["سمت سرور", "زبان Go با روتر chi", "کارایی بالا، تایپ ایستا، استقرار به‌صورت فایل اجرایی واحد"],
              ["پایگاه داده", "PostgreSQL", "پشتیبانی قوی از تراکنش و قیدهای یکپارچگی داده"],
              ["سمت کاربر", "React با TypeScript", "ساخت رابط تعاملی با ایمنی تایپ"],
              ["ظاهر رابط", "Tailwind CSS و قلم وزیرمتن", "پشتیبانی مناسب از چیدمان راست‌به‌چپ فارسی"],
              ["استقرار", "Docker Compose", "اجرای یکسان سه سرویس در محیط‌های مختلف"],
          ],
          widths=[2.8, 3.8, 8.4])

    heading(doc, "۵.۲. معماری لایه‌ای", 2)
    para(doc,
         "سامانه در چهار لایه سازمان یافته است و وابستگی‌ها تنها به سمت درون جریان دارند. "
         "این تفکیک باعث می‌شود هر نگرانی یک جایگاه مشخص داشته باشد:")
    code_block(doc, [
        "transport/http   ->   service/*   ->   storage/postgres",
        "                        domain",
    ])
    table(doc,
          ["لایه", "مسئولیت"],
          [
              ["دامنه", "موجودیت‌ها و مفاهیم کسب‌وکار به‌صورت خالص؛ بدون وابستگی به وب یا پایگاه داده"],
              ["سرویس", "منطق کسب‌وکار، تراکنش‌ها و تصمیم‌های دسترسی؛ هفت بسته به تفکیک حوزه"],
              ["ذخیره‌سازی", "پیاده‌سازی دسترسی به داده؛ تنها لایه‌ای که با کتابخانه نگاشت داده کار می‌کند"],
              ["انتقال", "مسیرهای وب، اعتبارسنجی ورودی، و نگاشت خطاها به کدهای وضعیت"],
          ],
          widths=[3.0, 12.0])
    para(doc,
         "دو نکته طراحی که به‌طور خاص رعایت شده‌اند: نخست، شِمای پایگاه داده تنها در "
         "اختیار فایل‌های مهاجرت است و از قابلیت ساخت خودکار جدول در کتابخانه نگاشت داده "
         "استفاده نکرده‌ام تا تغییرات شِما قابل بازبینی و قابل بازگشت باشند. دوم، ساختار "
         "داده‌ای که در پایگاه داده ذخیره می‌شود از ساختاری که در رابط برنامه‌نویسی منتشر "
         "می‌شود جدا است؛ بنابراین تغییر نام یک ستون، قرارداد بیرونی سامانه را نمی‌شکند.")

    heading(doc, "۵.۳. رابط برنامه‌نویسی و قرارداد آن", 2)
    para(doc,
         "رابط برنامه‌نویسی سامانه شامل سی‌ودو مسیر و چهل‌ودو عملیات است و به‌صورت رسمی در یک "
         "سند OpenAPI توصیف شده است. این سند مرجع قرارداد است: تایپ‌های سمت کاربر از آن "
         "تولید می‌شوند و دو آزمون خودکار، پاسخ‌های واقعی سرور را با آن تطبیق می‌دهند تا "
         "امکان واگرایی مستندات از پیاده‌سازی از بین برود.")
    para(doc,
         "درگاه یکپارچه‌سازی سامانه مادر جدا از نشست کاربران و با کلید دسترسی کار می‌کند و "
         "دو عملیات دارد: همگام‌سازی گروهی اطلاعات کارکنان بر اساس شماره پرسنلی، و استعلام "
         "وضعیت یک درخواست.")

    # ---------------------------------------------------- 6. implementation
    heading(doc, "۶. پیاده‌سازی", 1, new_page=True)

    heading(doc, "۶.۱. لاگ ممیزی", 2)
    para(doc,
         "تمام رویدادهای مؤثر ثبت می‌شوند: ورود به سامانه، ثبت و ارسال مجدد درخواست، شروع "
         "بررسی، تأیید، رد، بازگرداندن برای مدارک، بارگذاری مدرک، ثبت پرداخت، بستن "
         "پرونده، و تغییر پیکربندی قواعد. برای هر رویداد، کاربر، زمان و مقدار قبل و بعد "
         "نگه‌داری می‌شود.")
    para(doc,
         "نکته کلیدی پیاده‌سازی این است که نوشتن رکورد ممیزی و اعمال تغییر بر داده، در یک "
         "تراکنش پایگاه داده انجام می‌شوند. در نتیجه اگر هر یک از این دو با خطا مواجه شود، "
         "هر دو بازگردانده می‌شوند و از نظر ساختاری وضعیتی که تغییر یافته اما ردپا نداشته "
         "باشد قابل ایجاد نیست.")

    heading(doc, "۶.۲. احراز هویت و کنترل دسترسی", 2)
    bullets(doc, [
        "گذرواژه‌ها با الگوریتم bcrypt هش می‌شوند و هرگز به‌صورت متن آشکار ذخیره نمی‌گردند.",
        "احراز هویت کاربران با توکن امضاشده و اعتبار هشت‌ساعته انجام می‌شود.",
        "کنترل دسترسی در دو لایه اعمال می‌شود: نقش در سطح مسیر، و مالکیت داده در سطح سرویس.",
        "کلید درگاه یکپارچه‌سازی نیز تنها به‌صورت هش‌شده ذخیره می‌شود.",
        "سامانه در حالت تولید در صورتی که کلید امضای توکن مقدار پیش‌فرض توسعه باشد بالا نمی‌آید.",
    ])

    heading(doc, "۶.۳. دقت محاسبات مالی", 2)
    para(doc,
         "مبالغ در سراسر سامانه به‌صورت عدد صحیح ریال نگه‌داری و پردازش می‌شوند. دلیل این "
         "انتخاب آن است که نمایش اعشاری دودویی نمی‌تواند مقادیر ده‌دهی را به‌طور دقیق "
         "نگه دارد و این خطا در جمع‌های مالی و مقایسه با سقف‌ها انباشته می‌شود؛ ریال نیز در "
         "عمل واحد خردتر ندارد. بر این اساس، تمام محاسبه با عدد صحیح انجام می‌شود و تنها "
         "یک نقطه گردکردن در کل سامانه وجود دارد.")
    note(doc,
         "این تغییر را پس از تثبیت رفتار پیشین انجام دادم: نخست نوزده حالت مرزی محاسبه را "
         "در یک پرونده مرجع ثبت کردم، سپس تغییر را اعمال و نتیجه را مقایسه کردم. سیزده "
         "حالت بدون تغییر ماند و در شش حالت باقی‌مانده بیشترین اختلاف نیم ریال بود؛ یعنی "
         "تنها همان کسر زیر یک ریال حذف شد و تغییر معناداری در مبالغ رخ نداد.")

    heading(doc, "۶.۴. زمان کسب‌وکار", 2)
    para(doc,
         "دوره انتظار و سال قراردادی مفاهیمی بر پایه «روز تقویمی» هستند. اگر مرز روز با "
         "منطقه زمانی سرور محاسبه شود، یک درخواست نزدیک نیمه‌شب می‌تواند بر روی سرورهای "
         "مختلف نتیجه متفاوتی بدهد. به همین دلیل تمام تصمیم‌های مربوط به مرز روز در منطقه "
         "زمانی تهران ارزیابی می‌شوند و ساعت سامانه به‌صورت تزریق‌شده در اختیار سرویس‌ها "
         "قرار می‌گیرد تا رفتار آن قابل آزمون باشد.")

    heading(doc, "۶.۵. نگهداری مدارک پیوست", 2)
    para(doc,
         "مدارک، برخلاف سایر داده‌های سامانه، فایل‌اند و نه سطر پایگاه داده. فایل روی "
         "دیسک و در مسیری مشخص نگهداری می‌شود و تنها فراداده آن — نام اصلی فایل، زمان "
         "بارگذاری و کلید ذخیره‌سازی — در پایگاه داده ثبت می‌گردد. نام فایل روی دیسک را "
         "خود سامانه و به‌صورت شناسه یکتا تولید می‌کند؛ نام انتخابی کاربر تنها برای نمایش "
         "و دریافت به کار می‌رود و هرگز به مسیر تبدیل نمی‌شود؛ بنابراین نامی که نویسه‌های "
         "مسیر در آن جاسازی شده باشد نمی‌تواند به بیرون از پوشه ذخیره‌سازی دست پیدا کند.")
    para(doc,
         "نوع فایل از روی محتوای واقعی آن تشخیص داده می‌شود، نه از روی پسوند یا نوعی که "
         "مرورگر اعلام می‌کند؛ بنابراین تغییر پسوند یک فایل اجرایی به پسوند PDF آن را "
         "پذیرفتنی نمی‌کند. تنها چهار قالب PDF، JPEG، PNG و WebP و حداکثر تا پنج مگابایت "
         "پذیرفته می‌شوند.")
    para(doc,
         "ترتیب نوشتن اهمیت دارد: نخست فایل ذخیره می‌شود و سپس سطر فراداده و رکورد ممیزی "
         "در یک تراکنش ثبت می‌گردند؛ اگر آن تراکنش شکست بخورد فایل ذخیره‌شده حذف می‌شود. "
         "در نتیجه نه فایلی بدون سطر متناظر باقی می‌ماند و نه سطری که به فایل موجود اشاره "
         "نکند.")

    heading(doc, "۶.۶. رابط کاربری فارسی", 2)
    para(doc,
         "رابط کاربری شامل هفده صفحه و به‌طور کامل فارسی و راست‌به‌چپ است. تاریخ‌ها با "
         "تقویم هجری شمسی و اعداد با ارقام فارسی نمایش داده می‌شوند. قلم وزیرمتن همراه "
         "برنامه عرضه می‌شود تا نمایش مستقل از قلم‌های نصب‌شده روی دستگاه کاربر باشد. "
         "سه حالت نمایش روشن، تیره و خودکار پیش‌بینی شده است.")
    para(doc,
         "منوی کاربر بر اساس نقش او ساخته می‌شود و در صفحه جزئیات درخواست تنها دکمه‌هایی "
         "نمایش داده می‌شوند که در وضعیت جاری مجاز هستند؛ بنابراین کاربر با گزینه‌های "
         "نامعتبر مواجه نمی‌شود.")

    # -------------------------------------------------------------- 7. tests
    heading(doc, "۷. تضمین کیفیت و آزمون", 1, new_page=True)
    para(doc, "درستی سامانه در سه سطح راستی‌آزمایی شده است:")
    table(doc,
          ["سطح آزمون", "دامنه", "روش"],
          [
              ["آزمون واحد", "ریاضیات موتور قانون، تبدیل‌های مالی، پیکربندی", "بدون وابستگی بیرونی؛ اجرای سریع"],
              ["آزمون یکپارچگی", "سرویس‌ها روی پایگاه داده واقعی", "هر آزمون در یک تراکنش که در پایان بازگردانده می‌شود"],
              ["آزمون سرتاسری", "کل چرخه از رابط کاربری تا پایگاه داده", "هدایت خودکار مرورگر در چهارده گام"],
          ],
          widths=[3.2, 5.4, 6.4])
    para(doc,
         "آزمون سرتاسری، سناریوی کامل کسب‌وکار را می‌آزماید: ثبت درخواست توسط کارمند، "
         "بررسی و تأیید و پرداخت توسط کارشناس، مسیر رد با ذکر دلیل، انتشار نسخه جدید قانون "
         "توسط مدیر و راستی‌آزمایی این‌که درخواست بعدی با نرخ جدید محاسبه می‌شود، چرخه "
         "بازگرداندن برای مدارک همراه با بارگذاری فایل توسط کارمند و قفل‌شدن مدارک پس از "
         "ارسال مجدد، بررسی محدودیت‌های دسترسی، و صحت لاگ ممیزی و گزارش‌ها.")
    para(doc,
         "علاوه بر این، یک خط لوله یکپارچه‌سازی پیوسته تعریف شده است که در هر تغییر، "
         "بررسی سبک کد، ساخت پروژه و اجرای آزمون‌ها را به‌صورت خودکار انجام می‌دهد.")
    note(doc,
         "برای اطمینان از این‌که آزمون‌های تطبیق قرارداد واقعاً کار می‌کنند، به‌صورت عمدی "
         "یک اختلاف میان پیاده‌سازی و سند ایجاد کردم و تأیید کردم که آزمون آن را تشخیص "
         "می‌دهد و شکست می‌خورد.")

    # ------------------------------------------------------------ 8. results
    heading(doc, "۸. ارزیابی نتایج در برابر معیارهای پذیرش", 1)
    para(doc,
         "در پیشنهاد پروژه، پنج معیار پذیرش برای تعریف اتمام کار مشخص شده بود. وضعیت هر "
         "معیار به شرح زیر است:")
    table(doc,
          ["معیار پذیرش", "وضعیت", "شاهد"],
          [
              ["محاسبه درست مبلغ پرداختی و مانده سقف برای حداقل پنج نوع خدمت با درصد و سقف‌های متفاوت",
               "**محقق شد**", "آزمون‌های واحد موتور قانون و پرونده مرجع نوزده حالت مرزی"],
              ["اعمال تغییر قانون پوشش صرفاً از طریق پیکربندی و بدون تغییر کد",
               "**محقق شد**", "صفحه قوانین پوشش و گام اختصاصی در آزمون سرتاسری"],
              ["اجرای کامل گردش کار شامل مسیرهای تأیید، رد و بازگشت",
               "**محقق شد**", "آزمون‌های یکپارچگی گردش کار و آزمون سرتاسری"],
              ["ثبت کامل رویدادها در لاگ ممیزی و امکان بازسازی تاریخچه",
               "**محقق شد**", "صفحه تاریخچه اقدامات و آزمون ثبت رویدادهای چرخه کامل"],
              ["اعمال درست محدودیت دسترسی بر اساس نقش",
               "**محقق شد**", "آزمون‌های دسترسی در سطح سرویس و سرتاسری"],
          ],
          widths=[7.0, 2.4, 5.6])

    heading(doc, "۸.۱. اندازه پیاده‌سازی", 2)
    table(doc,
          ["شاخص", "مقدار", "شاخص ", "مقدار "],
          [
              ["کد سمت سرور", "حدود ۵٬۹۰۰ خط", "مسیرهای رابط برنامه‌نویسی", "۳۲ مسیر / ۴۲ عملیات"],
              ["کد آزمون سمت سرور", "حدود ۱٬۶۰۰ خط", "جداول پایگاه داده", "۱۲"],
              ["کد سمت کاربر", "حدود ۵٬۶۰۰ خط", "صفحات رابط کاربری", "۱۷"],
              ["نقش‌های کاربری", "۴", "وضعیت‌های درخواست", "۹"],
              ["انواع خدمت پیش‌فرض", "۵", "انواع رویداد ممیزی", "۱۱"],
          ],
          widths=[4.2, 3.3, 4.2, 3.3])

    # ------------------------------------------------- 9. limits & future work
    heading(doc, "۹. محدودیت‌ها و کارهای آینده", 1, new_page=True)
    para(doc,
         "علاوه بر مواردی که در بخش محدوده به‌عنوان خارج از دامنه پروژه معرفی شدند، "
         "موارد زیر برای استفاده عملیاتی در یک سازمان واقعی لازم‌اند و به‌عنوان ادامه کار "
         "پیشنهاد می‌شوند:")
    table(doc,
          ["مورد", "توضیح"],
          [
              ["سامانه اعلان", "اطلاع‌رسانی تغییر وضعیت درخواست از طریق رایانامه یا پیام کوتاه."],
              ["گزارش ماهانه بر پایه تقویم شمسی", "گروه‌بندی ماهانه در گزارش‌ها بر اساس ماه میلادی انجام می‌شود."],
              ["آزمون بار", "رفتار سامانه در حجم بالای داده و کاربر همزمان سنجیده نشده است."],
              ["احراز هویت دوعاملی", "با توجه به حساسیت داده‌های درمانی توصیه می‌شود."],
              ["گردش کار چندسطحی", "افزودن سطح تأیید مدیر در کنار کارشناس، برای سازمان‌هایی که چنین رویه‌ای دارند."],
              ["بارگذاری گروهی کارکنان", "برای ورود اولیه داده و انتقال از سامانه پیشین."],
          ],
          widths=[4.5, 10.5])
    para(doc,
         "همچنین از آن‌جا که هسته سامانه — یعنی ترکیب «سهمیه قابل پیکربندی» با «گردش کار "
         "چندمرحله‌ای» و «ممیزی کامل» — مستقل از حوزه بیمه است، می‌توان همان هسته را برای "
         "سایر خدمات رفاهی مانند وام، هزینه آموزش و هزینه سفر نیز به کار گرفت؛ این همان "
         "مسیری است که در پیشنهاد پروژه برای گسترش ماژول پیش‌بینی شده بود.")

    # ------------------------------------------------------------ 10. conclusion
    heading(doc, "۱۰. نتیجه‌گیری", 1)
    para(doc,
         "در این پروژه سامانه‌ای تحت وب برای مدیریت بیمه تکمیلی کارکنان طراحی و پیاده‌سازی "
         "کردم که سه مسئله اصلی روش دستی را هدف گرفته است: خطای محاسبه، نبود ردیابی، و "
         "هزینه بالای تغییر سیاست. محاسبه مبلغ قابل پرداخت و مانده سقف به‌صورت خودکار و بر "
         "پایه قواعد نسخه‌دار انجام می‌شود؛ گردش کار درخواست با مسیرهای تأیید، رد و بازگشت "
         "پیاده شده است؛ و تمام رویدادهای مؤثر به‌صورت تفکیک‌ناپذیر از تغییر داده در لاگ "
         "ممیزی ثبت می‌شوند.")
    para(doc,
         "دستاوردی که بیش از همه بر آن تأکید دارم، جدا کردن قواعد کسب‌وکار از کد است. با "
         "این طراحی، تغییر سیاست رفاهی سازمان به یک عملیات داده‌ای تبدیل می‌شود که مدیر "
         "سامانه خودش انجام می‌دهد و همان لحظه اعمال می‌گردد، در حالی که خودِ این تغییر نیز "
         "ممیزی شده و قابل بازبینی است.")
    para(doc,
         "پنج معیار پذیرشی که در پیشنهاد پروژه تعریف شده بود، محقق شده و با آزمون‌های "
         "خودکار در سه سطح راستی‌آزمایی شده‌اند. محدودیت‌های باقی‌مانده نیز در بخش پیشین "
         "به‌صورت صریح فهرست شده‌اند تا مسیر تکمیل سامانه برای استفاده عملیاتی روشن باشد.")

    # ------------------------------------------------------------- appendix
    heading(doc, "پیوست الف — راهنمای اجرای سامانه", 1, new_page=True)
    para(doc, "سامانه از سه سرویس تشکیل شده است: پایگاه داده، سرویس برنامه، و رابط کاربری. "
              "برای اجرا با استفاده از Docker:")
    # Persian must not go inside a monospace block: code fonts carry no Persian
    # glyphs, so the letters render unjoined. Command and explanation are split
    # into a table instead.
    table(doc,
          ["فرمان", "کارکرد"],
          [
              ["make up", "ساخت تصویرها و اجرای هر سه سرویس"],
              ["make seed", "بارگذاری داده نمونه"],
              ["make down", "توقف سرویس‌ها"],
          ],
          widths=[4.0, 11.0])
    para(doc, "پس از اجرا، رابط کاربری روی نشانی زیر در دسترس است:")
    code_block(doc, ["http://localhost:5173"])
    para(doc, "حساب‌های نمونه برای بررسی نقش‌های مختلف:")
    table(doc,
          ["نام کاربری", "گذرواژه", "نقش"],
          [
              ["admin", "Admin123!", "مدیر سامانه"],
              ["reviewer", "Reviewer123!", "کارشناس بررسی"],
              ["sara.ahmadi", "Employee123!", "کارمند — طرح استاندارد"],
              ["reza.karimi", "Employee123!", "کارمند — طرح ویژه"],
              ["auditor", "Auditor123!", "ممیز"],
          ],
          widths=[4.0, 4.0, 7.0])
    note(doc, "گذرواژه‌های بالا صرفاً برای داده نمونه و محیط آزمایش است و در استقرار واقعی "
              "باید تغییر کنند. سرویس برنامه در حالت تولید با کلید امضای پیش‌فرض بالا نمی‌آید.")

    heading(doc, "پیوست ب — ساختار پوشه‌های پروژه", 1)
    table(doc,
          ["مسیر", "محتوا"],
          [
              ["backend/api/openapi.yaml", "سند رسمی رابط برنامه‌نویسی"],
              ["backend/cmd/api", "نقطه ورود سرویس برنامه"],
              ["backend/cmd/seed", "بارگذاری داده نمونه"],
              ["backend/internal/domain", "موجودیت‌ها و مفاهیم کسب‌وکار"],
              ["backend/internal/service", "منطق کسب‌وکار (هفت بسته)"],
              ["backend/internal/storage", "دسترسی به پایگاه داده"],
              ["backend/internal/transport", "مسیرهای وب و نگاشت خطا"],
              ["backend/migrations", "مهاجرت‌های شِمای پایگاه داده"],
              ["frontend/src/api", "لایه ارتباط با سرویس برنامه"],
              ["frontend/src/pages", "صفحات رابط کاربری"],
              ["frontend/src/components", "اجزای مشترک رابط کاربری"],
              ["e2e/", "آزمون سرتاسری روی مرورگر"],
              ["deploy/", "تنظیمات استقرار"],
              ["docs/", "مستندات پروژه"],
          ],
          widths=[6.0, 9.0])

    heading(doc, "پیوست ج — فهرست مستندات پروژه", 1)
    table(doc,
          ["سند", "محتوا"],
          [
              ["ARCHITECTURE.md", "معماری سامانه و تصمیم‌های طراحی"],
              ["ERD.md", "نمودار موجودیت–رابطه و شرح جداول"],
              ["USE-CASES.md", "سناریوهای کاربری و جریان‌های جایگزین"],
              ["API-CONTRACT.md", "شرح رابط برنامه‌نویسی"],
              ["openapi.yaml", "توصیف رسمی و ماشین‌خوان رابط برنامه‌نویسی"],
              ["adr/", "مستند تصمیم‌های معماری"],
          ],
          widths=[4.5, 10.5])

    return doc


if __name__ == "__main__":
    import pathlib
    out = pathlib.Path(__file__).parent / "گزارش-پروژه-پایانی-بیمه-تکمیلی.docx"
    build().save(out)
    print(f"ساخته شد: {out}")
